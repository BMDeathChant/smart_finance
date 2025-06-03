"""Excel数据分析测试脚本
使用openai_wrapper中的DataAnalyzer分析Excel文件
"""

import pandas as pd
from openai_wrapper import DataAnalyzer, DeepseekWrapper

def analyze_excel_file(excel_path: str, api_key: str):
    """分析Excel文件并生成报告
    
    Args:
        excel_path: Excel文件路径
        api_key: Deepseek API密钥
    """
    from docx import Document
    from datetime import datetime
    import os
    
    # 初始化Deepseek客户端
    wrapper = DeepseekWrapper(api_key)
    analyzer = DataAnalyzer(wrapper)
    
    # 创建Word文档
    doc = Document()
    doc.add_heading('Excel数据分析报告', level=1)
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"源文件: {os.path.basename(excel_path)}")
    
    # 读取Excel文件
    with pd.ExcelFile(excel_path) as xls:
        # 分析每个sheet
        for sheet_name in xls.sheet_names:
            print(f"\n=== 正在分析工作表: {sheet_name} ===")
            doc.add_heading(f"工作表: {sheet_name}", level=2)
            
            df = pd.read_excel(xls, sheet_name)
            
            # 执行分析
            result = analyzer.analyze(df)
            
            # 控制台输出
            print("\n分析摘要:")
            print(result.summary)
            print("\n原始响应:")
            print(result.raw_response)
            
            # 写入Word文档
            doc.add_heading('分析摘要', level=3)
            doc.add_paragraph(result.summary)
            doc.add_heading('详细分析', level=3)
            doc.add_paragraph(result.raw_response)
    
    # 保存Word文档
    report_path = os.path.join(
        os.path.dirname(excel_path),
        f"分析报告_{os.path.splitext(os.path.basename(excel_path))[0]}.docx"
    )
    doc.save(report_path)
    print(f"\n分析报告已保存至: {report_path}")

if __name__ == "__main__":
    import tkinter as tk
    from tkinter import filedialog
    
    # 创建文件选择对话框
    root = tk.Tk()
    root.withdraw()  # 隐藏主窗口
    
    # 选择Excel文件
    excel_path = filedialog.askopenfilename(
        title="选择Excel文件",
        filetypes=[("Excel文件", "*.xlsx *.xls"), ("所有文件", "*.*")]
    )
    
    if not excel_path:
        print("未选择文件")
        exit()
    
    # 输入API密钥
    api_key = input("请输入Deepseek API密钥: ").strip()
    if not api_key:
        print("API密钥不能为空")
        exit()
    
    # 执行分析
    analyze_excel_file(excel_path, api_key)
